"""
Извлечение текста из DOCX файлов
"""

from docx import Document as DocxDocument


def extract_text_from_docx(file_path):
    """
    Извлекает текст из DOCX файла
    
    Args:
        file_path: путь к .docx файлу
        
    Returns:
        str: извлечённый текст
    """
    try:
        doc = DocxDocument(file_path)
        
        # Извлекаем текст из параграфов
        paragraphs_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs_text.append(text)
        
        # Извлекаем текст из таблиц
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    tables_text.append(' | '.join(row_text))
        
        # Объединяем весь текст
        full_text = '\n'.join(paragraphs_text)
        
        if tables_text:
            full_text += '\n\n[ТАБЛИЦЫ]\n' + '\n'.join(tables_text)
        
        return full_text
        
    except Exception as e:
        raise Exception(f"Ошибка при извлечении текста из DOCX: {str(e)}")
